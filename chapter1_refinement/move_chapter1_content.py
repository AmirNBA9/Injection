from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def find_paragraph_range(doc, start_text, end_text=None):
    """پیدا کردن محدوده پاراگراف‌ها بر اساس متن"""
    start_idx = None
    end_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if start_text in text and start_idx is None:
            start_idx = i
        if end_text and end_text in text and start_idx is not None:
            end_idx = i
            break
    
    return start_idx, end_idx

def find_insertion_point(doc, marker_text, after=True):
    """پیدا کردن نقطه درج بر اساس متن مارکر"""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if marker_text in text:
            return i + 1 if after else i
    return None

def copy_paragraphs(source_doc, start_idx, end_idx):
    """کپی پاراگراف‌ها از یک سند"""
    copied = []
    for i in range(start_idx, end_idx + 1):
        if i < len(source_doc.paragraphs):
            para = source_doc.paragraphs[i]
            # کپی محتوای پاراگراف
            copied.append({
                'text': para.text,
                'style': para.style.name if para.style else None,
                'runs': [{'text': run.text, 'bold': run.bold, 'italic': run.italic} for run in para.runs]
            })
    return copied

def insert_paragraphs(doc, target_idx, paragraphs_data, add_marker=True):
    """درج پاراگراف‌ها در موقعیت مشخص"""
    if add_marker:
        # اضافه کردن پاراگراف علامت‌گذاری
        marker_para = doc.paragraphs[target_idx].insert_paragraph_before()
        marker_para.add_run("[منتقل‌شده از فصل ۱]").bold = True
        marker_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        target_idx += 1
    
    for para_data in paragraphs_data:
        new_para = doc.paragraphs[target_idx].insert_paragraph_before()
        new_para.text = para_data['text']
        if para_data['style']:
            try:
                new_para.style = para_data['style']
            except:
                pass
        
        # بازسازی formatting
        if para_data['runs']:
            new_para.clear()
            for run_data in para_data['runs']:
                run = new_para.add_run(run_data['text'])
                if run_data.get('bold'):
                    run.bold = True
                if run_data.get('italic'):
                    run.italic = True
        
        target_idx += 1
    
    return target_idx

def remove_paragraphs(doc, start_idx, count):
    """حذف پاراگراف‌ها از موقعیت مشخص"""
    for i in range(count):
        if start_idx < len(doc.paragraphs):
            para = doc.paragraphs[start_idx]
            para._element.getparent().remove(para._element)

# باز کردن سند
doc = Document(r'd:\Workingdir\Injection\Thesis_V3.3.docx')

print("Starting content transfer...")

# انتقال ۶: بخش 1-2-2 (P192-P196) → انتهای 2-1-1
print("\n1. Transferring section 1-2-2 to end of 2-1-1...")
# پیدا کردن مبدا
src_start_6 = 192  # P192
src_end_6 = 196    # P196
# پیدا کردن مقصد (بعد از آخرین پاراگراف 2-1-1، قبل از 2-1-2)
dest_6 = None
for i in range(434, 442):
    para = doc.paragraphs[i]
    if '2-1-2' in para.text:
        dest_6 = i
        break

if dest_6:
    copied_6 = copy_paragraphs(doc, src_start_6, src_end_6)
    insert_paragraphs(doc, dest_6, copied_6, add_marker=True)
    # حذف از مبدا (بعد از درج)
    remove_paragraphs(doc, src_start_6, src_end_6 - src_start_6 + 1)
    print(f"   ✓ Transferred paragraphs {src_start_6}-{src_end_6} to position {dest_6}")

# انتقال ۵: بخش 1-2-3 (P197-P200) → انتهای 2-1-1
print("\n2. Transferring section 1-2-3 to end of 2-1-1...")
src_start_5 = 197
src_end_5 = 200
# مقصد همان جایی است که انتقال ۶ انجام شد (بعد از آن)
dest_5 = dest_6 + len(copied_6) + 1 if dest_6 else None

if dest_5:
    copied_5 = copy_paragraphs(doc, src_start_5, src_end_5)
    insert_paragraphs(doc, dest_5, copied_5, add_marker=True)
    remove_paragraphs(doc, src_start_5, src_end_5 - src_start_5 + 1)
    print(f"   ✓ Transferred paragraphs {src_start_5}-{src_end_5} to position {dest_5}")

# انتقال ۴: بخش 1-2-1 (P177-P191) → انتهای 2-1-2
print("\n3. Transferring section 1-2-1 to end of 2-1-2...")
src_start_4 = 177
src_end_4 = 191
# پیدا کردن انتهای 2-1-2 (قبل از 2-1-3)
dest_4 = None
for i in range(442, 470):
    para = doc.paragraphs[i]
    if '2-1-3' in para.text:
        dest_4 = i
        break

if dest_4:
    copied_4 = copy_paragraphs(doc, src_start_4, src_end_4)
    insert_paragraphs(doc, dest_4, copied_4, add_marker=True)
    remove_paragraphs(doc, src_start_4, src_end_4 - src_start_4 + 1)
    print(f"   ✓ Transferred paragraphs {src_start_4}-{src_end_4} to position {dest_4}")

# انتقال ۳: بخش 1-2-4 (P201-P224) → انتهای 2-1-2
print("\n4. Transferring section 1-2-4 to end of 2-1-2...")
src_start_3 = 201
src_end_3 = 224
# مقصد بعد از انتقال ۴
dest_3 = dest_4 + len(copied_4) + 1 if dest_4 else None

if dest_3:
    copied_3 = copy_paragraphs(doc, src_start_3, src_end_3)
    insert_paragraphs(doc, dest_3, copied_3, add_marker=True)
    remove_paragraphs(doc, src_start_3, src_end_3 - src_start_3 + 1)
    print(f"   ✓ Transferred paragraphs {src_start_3}-{src_end_3} to position {dest_3}")

# انتقال ۲: بخش 1-2-5 (P225-P256) → ابتدای فصل ۳
print("\n5. Transferring section 1-2-5 to beginning of Chapter 3...")
src_start_2 = 225
src_end_2 = 256
# پیدا کردن ابتدای فصل ۳ (بعد از مقدمه، قبل از 3-1)
dest_2 = None
for i in range(823, 860):
    para = doc.paragraphs[i]
    if '3-1' in para.text and 'Heading' in para.style.name:
        dest_2 = i
        break

if dest_2:
    copied_2 = copy_paragraphs(doc, src_start_2, src_end_2)
    insert_paragraphs(doc, dest_2, copied_2, add_marker=True)
    remove_paragraphs(doc, src_start_2, src_end_2 - src_start_2 + 1)
    print(f"   ✓ Transferred paragraphs {src_start_2}-{src_end_2} to position {dest_2}")

# انتقال ۱: بخش 1-4 (P276-P296) → انتهای 2-4
print("\n6. Transferring section 1-4 to end of 2-4...")
src_start_1 = 276
src_end_1 = 296
# پیدا کردن انتهای 2-4 (بعد از 2-4-3)
dest_1 = None
for i in range(589, 823):
    para = doc.paragraphs[i]
    if 'فصل سوم' in para.text or 'فصل ۳' in para.text:
        dest_1 = i
        break

if dest_1:
    copied_1 = copy_paragraphs(doc, src_start_1, src_end_1)
    insert_paragraphs(doc, dest_1, copied_1, add_marker=True)
    remove_paragraphs(doc, src_start_1, src_end_1 - src_start_1 + 1)
    print(f"   ✓ Transferred paragraphs {src_start_1}-{src_end_1} to position {dest_1}")

# ذخیره فایل
doc.save(r'd:\Workingdir\Injection\Thesis_V3.3.docx')
print("\n✓ All transfers completed successfully!")
print("✓ File saved as Thesis_V3.3.docx")
