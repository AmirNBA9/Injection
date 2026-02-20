from docx import Document

doc = Document(r'd:\Workingdir\Injection\Thesis_V3.4.docx')

# پیدا کردن فصل اول
ch1_start = None
ch1_end = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if 'فصل اول' in text and 'Heading' in style:
        ch1_start = i
    if 'فصل دوم' in text and 'Heading' in style:
        ch1_end = i
        break

print(f"Chapter 1: P{ch1_start} to P{ch1_end}")
print(f"\nTotal paragraphs in Chapter 1: {ch1_end - ch1_start}")

# بررسی بخش‌های مختلف
print("\n=== Section 1-2 ===")
for i in range(ch1_start, min(ch1_start + 40, ch1_end, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if '1-2' in text or (i > ch1_start + 20 and len(text) > 50):
        print(f"P{i}: {text[:120]}")

print("\n=== Section 1-6 ===")
for i in range(ch1_start, min(ch1_end, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if '1-6' in text:
        # نمایش چند پاراگراف بعد
        for j in range(i, min(i + 5, ch1_end, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            if p.text.strip():
                print(f"P{j}: {p.text[:120]}")

print("\n=== Section 1-7 ===")
for i in range(ch1_start, min(ch1_end, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if '1-7' in text:
        # نمایش چند پاراگراف بعد
        for j in range(i, min(i + 10, ch1_end, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            if p.text.strip():
                print(f"P{j}: {p.text[:120]}")

print("\n✓ Check complete!")
