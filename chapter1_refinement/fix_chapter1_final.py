from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document(r'd:\Workingdir\Injection\Thesis_V3.4.docx')

print("Fixing Chapter 1 issues...\n")

# پیدا کردن بخش‌ها
ch1_start = None
sections = {}

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if 'فصل اول' in text and 'Heading' in style:
        ch1_start = i
    
    if ch1_start is not None and 'Heading' in style:
        if '1-2' in text and '1-2-' not in text:
            sections['1-2'] = i
        elif '1-3' in text:
            sections['1-3'] = i
        elif '1-6' in text:
            sections['1-6'] = i
        elif '1-7' in text:
            sections['1-7'] = i

print(f"Found sections: {sections}")

# اصلاح بخش 1-2: حذف محتوای معکوس و نوشتن مجدد
if '1-2' in sections:
    print("\nFixing section 1-2...")
    sec_1_2_start = sections['1-2']
    sec_1_2_end = sections.get('1-3', sec_1_2_start + 20)
    
    # حذف همه پاراگراف‌های بعد از عنوان تا بخش بعدی
    indices_to_remove = []
    for i in range(sec_1_2_start + 1, min(sec_1_2_end, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.text.strip():
            # اگر به بخش بعدی رسیدیم، توقف
            if '1-3' in para.text and 'Heading' in (para.style.name if para.style else ''):
                break
            indices_to_remove.append(i)
    
    # حذف از آخر به اول
    for i in reversed(indices_to_remove):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            para._element.getparent().remove(para._element)
    
    # پیدا کردن پاراگراف بعدی برای درج
    next_idx = sec_1_2_start + 1
    while next_idx < len(doc.paragraphs) and not doc.paragraphs[next_idx].text.strip():
        next_idx += 1
    
    # نوشتن متن جدید به ترتیب صحیح
    para1 = doc.paragraphs[next_idx].insert_paragraph_before()
    para1.add_run("""این پژوهش با هدف ارائه راهکاری جامع برای تشخیص حملات تزریق پایگاه‌داده غیررابطه‌ای طراحی شده است. برای دست‌یابی به این هدف، مجموعه‌داده‌های اولیه به‌صورت دستی تحلیل و برچسب‌گذاری می‌شوند تا کیفیت داده‌های پایه تضمین گردد. سپس از مدل‌های زبان بزرگ برای تولید داده‌های مصنوعی متنوع و واقعی‌تر استفاده می‌شود که به غنی‌سازی مجموعه‌دادگان کمک می‌کند. در مرحله بعد، تکنیک‌های مختلف یادگیری ماشینی برای ارزیابی و دسته‌بندی داده‌های موجود و مصنوعی به‌کار گرفته می‌شوند.""")
    
    para2 = doc.paragraphs[next_idx].insert_paragraph_before()
    para2.add_run("""مفاهیم کلیدی این پژوهش شامل شناسایی و مقابله با حملات تزریق پایگاه‌داده غیررابطه‌ای است که از نقاط ضعف این نوع پایگاه‌ها سوءاستفاده می‌کنند. این حملات معمولاً از طریق ارسال درخواست‌های مخرب و دست‌کاری پرس‌وجوها انجام می‌گیرند. همچنین، تحلیل و طراحی مجموعه‌دادگان شامل داده‌های اولیه و داده‌های مصنوعی تولیدشده با مدل‌های زبان بزرگ، یکی از محورهای اصلی پژوهش است. در این راستا، به‌جای استفاده از شبکه‌های مولد متخاصم، از مدل‌های زبان بهره گرفته می‌شود که توانایی تولید داده‌های متنی و ساختاریافته باکیفیت بالا را دارند.""")
    
    para3 = doc.paragraphs[next_idx].insert_paragraph_before()
    para3.add_run("""این پژوهش شامل طراحی و ارزیابی یک سیستم یادگیری گروهی است که از مدل‌های مختلف مانند رگرسیون لجستیک، جنگل تصادفی و XGBoost تشکیل شده است. این مدل‌ها با ترکیب قدرت یادگیری فردی خود، دقت شناسایی حملات را بهبود می‌بخشند. دستاوردهای مورد انتظار این پژوهش شامل ارائه روشی نوین برای شناسایی حملات تزریق پایگاه‌داده غیررابطه‌ای، ایجاد مجموعه‌دادگان جامع و متنوع، و بهبود دقت مدل‌های یادگیری ماشینی در محیط‌های واقعی است.""")
    
    print("  ✓ Section 1-2 fixed")

# بازنویسی کامل بخش 1-6
if '1-6' in sections:
    print("\nRewriting section 1-6...")
    sec_1_6_start = sections['1-6']
    sec_1_6_end = sections.get('1-7', sec_1_6_start + 30)
    
    # حذف همه محتوای قدیمی
    indices_to_remove = []
    for i in range(sec_1_6_start + 1, min(sec_1_6_end, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.text.strip():
            if '1-7' in para.text and 'Heading' in (para.style.name if para.style else ''):
                break
            indices_to_remove.append(i)
    
    for i in reversed(indices_to_remove):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            para._element.getparent().remove(para._element)
    
    # پیدا کردن پاراگراف بعدی
    next_idx = sec_1_6_start + 1
    while next_idx < len(doc.paragraphs) and not doc.paragraphs[next_idx].text.strip():
        next_idx += 1
    
    # نوشتن متن جدید خلاصه و متمرکز
    para1 = doc.paragraphs[next_idx].insert_paragraph_before()
    para1.add_run("""نوآوری اصلی این پژوهش در تمرکز بر حملات تزریق پایگاه‌داده غیررابطه‌ای است که به‌دلیل تفاوت‌های ساختاری اساسی با پایگاه‌داده‌های رابطه‌ای، نیازمند رویکردهای امنیتی متفاوت و پیچیده‌تری هستند. برخلاف پایگاه‌داده‌های رابطه‌ای که از ساختار جدولی و زبان پرس‌وجوی استاندارد SQL استفاده می‌کنند، پایگاه‌داده‌های غیررابطه‌ای مانند MongoDB از ساختارهای غیرجدولی و زبان‌های پرس‌وجوی خاص خود بهره می‌برند. این تفاوت‌ها باعث می‌شود که حملات تزریق در این نوع پایگاه‌ها از الگوهای پیچیده‌تر و غیرمعمول‌تری استفاده کنند که شناسایی آن‌ها با روش‌های سنتی مبتنی بر SQL دشوار است.""")
    
    para2 = doc.paragraphs[next_idx].insert_paragraph_before()
    para2.add_run("""در این پژوهش، با استفاده از مدل‌های زبان بزرگ برای تولید داده‌های مصنوعی مرتبط با حملات MongoDB، یک مجموعه‌داده جامع و متنوع ایجاد شده است. همچنین، یک سیستم یادگیری گروهی طراحی شده که از ترکیب چندین مدل یادگیری ماشینی برای بهبود دقت تشخیص استفاده می‌کند. این رویکرد جامع، خلأ موجود در پژوهش‌های امنیت سایبری را پر کرده و گامی مهم در راستای ارتقای امنیت پایگاه‌داده‌های غیررابطه‌ای برداشته است.""")
    
    print("  ✓ Section 1-6 rewritten")

# بازنویسی بخش 1-7
if '1-7' in sections:
    print("\nRewriting section 1-7...")
    sec_1_7_start = sections['1-7']
    
    # پیدا کردن انتهای فصل اول
    ch1_end = None
    for i in range(sec_1_7_start + 1, min(sec_1_7_start + 100, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if 'فصل' in para.text and 'Heading' in (para.style.name if para.style else ''):
            ch1_end = i
            break
    
    if ch1_end is None:
        ch1_end = sec_1_7_start + 50
    
    # حذف محتوای قدیمی
    indices_to_remove = []
    for i in range(sec_1_7_start + 1, min(ch1_end, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.text.strip():
            indices_to_remove.append(i)
    
    for i in reversed(indices_to_remove):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            para._element.getparent().remove(para._element)
    
    # پیدا کردن پاراگراف بعدی
    next_idx = sec_1_7_start + 1
    while next_idx < len(doc.paragraphs) and not doc.paragraphs[next_idx].text.strip():
        next_idx += 1
    
    # نوشتن پاراگراف‌های جدید
    para1 = doc.paragraphs[next_idx].insert_paragraph_before()
    para1.add_run("""این پایان‌نامه در شش فصل تنظیم شده است و مسیر پژوهش را از مبانی نظری تا جمع‌بندی و پیشنهادهای آتی به‌صورت منسجم دنبال می‌کند.""")
    
    para2 = doc.paragraphs[next_idx].insert_paragraph_before()
    para2.add_run("""فصل دوم به بررسی مبانی نظری و پیشینه پژوهش اختصاص دارد. در این فصل، مفاهیم پایه مرتبط با حملات تزریق در پایگاه‌داده‌های غیررابطه‌ای، کاربردهای یادگیری ماشین در امنیت، و رویکردهای تولید داده مصنوعی با مدل‌های زبان بزرگ در مقایسه با روش‌های سنتی مرور می‌شوند. همچنین، شکاف‌های پژوهشی و تمایزات رویکرد پیشنهادی جمع‌بندی می‌گردد.""")
    
    para3 = doc.paragraphs[next_idx].insert_paragraph_before()
    para3.add_run("""فصل سوم روش پیشنهادی را تشریح می‌کند. این فصل شامل چارچوب کلی پژوهش، فرایند تولید داده مصنوعی مبتنی بر مدل‌های زبان بزرگ، معرفی مدل‌های یادگیری ماشین و مدل‌های گروهی به‌کاررفته و منطق انتخاب آن‌ها، تنظیمات کلیدی و پارامترها، و معماری پیاده‌سازی سیستم پیشنهادی است.""")
    
    para4 = doc.paragraphs[next_idx].insert_paragraph_before()
    para4.add_run("""فصل چهارم به تحلیل نتایج و بحث اختصاص دارد. در این فصل، مراحل آماده‌سازی داده‌ها، اجرای مدل‌ها و نحوه ارزیابی بیان می‌شود. سپس عملکرد مدل‌ها با معیارهایی مانند دقت، شاخص‌های خطا، و معیارهای زمانی و کارایی مقایسه و تحلیل می‌گردد. همچنین، تفسیر یافته‌ها، مقایسه با پیشینه پژوهش، بررسی نقاط قوت و محدودیت‌ها، و تحلیل اثر داده‌های مصنوعی و انتخاب مدل‌ها بر عملکرد و تعمیم‌پذیری ارائه می‌شود.""")
    
    para5 = doc.paragraphs[next_idx].insert_paragraph_before()
    para5.add_run("""فصل پنجم شامل جمع‌بندی و کارهای آتی است. در این فصل، دستاوردهای اصلی پژوهش، مرور نوآوری‌ها، بیان محدودیت‌ها و ارائه پیشنهادهای مشخص برای توسعه و پژوهش‌های آینده ارائه می‌شود.""")
    
    print("  ✓ Section 1-7 rewritten")

# ذخیره فایل
doc.save(r'd:\Workingdir\Injection\Thesis_V3.4.docx')
print("\n✓ All fixes applied successfully!")
print("✓ File saved as Thesis_V3.4.docx")
