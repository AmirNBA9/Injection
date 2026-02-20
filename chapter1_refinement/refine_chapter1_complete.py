from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document(r'd:\Workingdir\Injection\Thesis_V3.4.docx')

print("Starting Chapter 1 refinement...")

# پیدا کردن محدوده فصل اول
ch1_start = None
ch1_end = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if 'فصل اول' in text and 'Heading' in style:
        ch1_start = i
    if 'فصل دوم' in text and 'Heading' in style:
        ch1_end = i
        break

print(f"Chapter 1: P{ch1_start} to P{ch1_end}")

# پیدا کردن بخش‌های مختلف
sections = {}
for i in range(ch1_start, min(ch1_end, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if 'Heading' in style:
        if '1-1' in text:
            sections['1-1'] = i
        elif '1-2' in text and '1-2-' not in text:  # فقط 1-2 نه 1-2-1
            sections['1-2'] = i
        elif '1-3' in text:
            sections['1-3'] = i
        elif '1-5' in text:
            sections['1-5'] = i
        elif '1-6' in text:
            sections['1-6'] = i
        elif '1-7' in text:
            sections['1-7'] = i

print(f"Found sections: {sections}")

# مرحله 1: حذف بخش‌های منتقل‌شده (1-2-1, 1-2-2, 1-2-3, 1-2-4, 1-2-5, 1-4)
print("\nStep 1: Removing transferred sections...")
removed_sections = ['1-2-1', '1-2-2', '1-2-3', '1-2-4', '1-2-5', '1-4']
indices_to_remove = set()

# پیدا کردن بخش‌های منتقل‌شده
for i in range(ch1_start, min(ch1_end, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    # پیدا کردن عنوان‌های بخش‌های منتقل‌شده
    if 'Heading' in style:
        for sec in removed_sections:
            if sec in text:
                # پیدا کردن محدوده این بخش تا بخش بعدی
                start_idx = i
                end_idx = None
                for j in range(i + 1, min(ch1_end + 50, len(doc.paragraphs))):  # کمی بیشتر جستجو کنیم
                    if j >= len(doc.paragraphs):
                        break
                    next_para = doc.paragraphs[j]
                    next_text = next_para.text.strip()
                    next_style = next_para.style.name if next_para.style else ''
                    if 'Heading' in next_style:
                        # اگر بخش بعدی پیدا شد
                        if any(x in next_text for x in ['1-1', '1-2', '1-3', '1-5', '1-6', '1-7', 'فصل']):
                            # بررسی کنیم که آیا این بخش بعدی است یا نه
                            if not any(removed_sec in next_text for removed_sec in removed_sections):
                                end_idx = j
                                break
                if end_idx is None:
                    end_idx = min(ch1_end, len(doc.paragraphs))
                
                # اضافه کردن به لیست حذف
                for k in range(start_idx, end_idx):
                    indices_to_remove.add(k)
                print(f"  Found section {sec}: P{start_idx}-P{end_idx}")

# حذف از آخر به اول
removed_count = 0
for i in sorted(indices_to_remove, reverse=True):
    if i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        para._element.getparent().remove(para._element)
        removed_count += 1

print(f"  ✓ Removed {removed_count} paragraphs")

# پیدا کردن مجدد بخش‌ها بعد از حذف
sections = {}
ch1_end_new = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if 'فصل اول' in text and 'Heading' in style:
        ch1_start = i
    if 'فصل دوم' in text and 'Heading' in style:
        ch1_end_new = i
        break

for i in range(ch1_start, min(ch1_end_new if ch1_end_new else ch1_end, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if 'Heading' in style:
        if '1-1' in text:
            sections['1-1'] = i
        elif '1-2' in text and '1-2-' not in text:
            sections['1-2'] = i
        elif '1-3' in text:
            sections['1-3'] = i
        elif '1-5' in text:
            sections['1-5'] = i
        elif '1-6' in text:
            sections['1-6'] = i
        elif '1-7' in text:
            sections['1-7'] = i

if ch1_end_new:
    ch1_end = ch1_end_new

print(f"Updated sections after removal: {sections}")

# مرحله 2: بازنویسی بخش 1-2 به متن روان
print("\nStep 2: Rewriting section 1-2 as flowing prose...")

# پیدا کردن محدوده بخش 1-2
if '1-2' in sections:
    sec_1_2_start = sections['1-2']
    sec_1_2_end = sections.get('1-3', ch1_end)
    
    # خواندن محتوای فعلی
    current_content = []
    for i in range(sec_1_2_start, min(sec_1_2_end, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.text.strip():
            current_content.append(para.text.strip())
    
    # حذف محتوای قدیمی (بعد از عنوان)
    for i in range(sec_1_2_start + 1, min(sec_1_2_end, len(doc.paragraphs))):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            if para.text.strip():
                para._element.getparent().remove(para._element)
    
    # پیدا کردن پاراگراف بعدی برای درج
    next_para_idx = sec_1_2_start + 1
    while next_para_idx < len(doc.paragraphs) and not doc.paragraphs[next_para_idx].text.strip():
        next_para_idx += 1
    
    # نوشتن متن جدید روان
    new_para_1 = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para_1.add_run("""این پژوهش با هدف ارائه راهکاری جامع برای تشخیص حملات تزریق پایگاه‌داده غیررابطه‌ای طراحی شده است. برای دست‌یابی به این هدف، مجموعه‌داده‌های اولیه به‌صورت دستی تحلیل و برچسب‌گذاری می‌شوند تا کیفیت داده‌های پایه تضمین گردد. سپس از مدل‌های زبان بزرگ برای تولید داده‌های مصنوعی متنوع و واقعی‌تر استفاده می‌شود که به غنی‌سازی مجموعه‌دادگان کمک می‌کند. در مرحله بعد، تکنیک‌های مختلف یادگیری ماشینی برای ارزیابی و دسته‌بندی داده‌های موجود و مصنوعی به‌کار گرفته می‌شوند.""")
    
    new_para_2 = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para_2.add_run("""مفاهیم کلیدی این پژوهش شامل شناسایی و مقابله با حملات تزریق پایگاه‌داده غیررابطه‌ای است که از نقاط ضعف این نوع پایگاه‌ها سوءاستفاده می‌کنند. این حملات معمولاً از طریق ارسال درخواست‌های مخرب و دست‌کاری پرس‌وجوها انجام می‌گیرند. همچنین، تحلیل و طراحی مجموعه‌دادگان شامل داده‌های اولیه و داده‌های مصنوعی تولیدشده با مدل‌های زبان بزرگ، یکی از محورهای اصلی پژوهش است. در این راستا، به‌جای استفاده از شبکه‌های مولد متخاصم، از مدل‌های زبان بهره گرفته می‌شود که توانایی تولید داده‌های متنی و ساختاریافته باکیفیت بالا را دارند.""")
    
    new_para_3 = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para_3.add_run("""این پژوهش شامل طراحی و ارزیابی یک سیستم یادگیری گروهی است که از مدل‌های مختلف مانند رگرسیون لجستیک، جنگل تصادفی و XGBoost تشکیل شده است. این مدل‌ها با ترکیب قدرت یادگیری فردی خود، دقت شناسایی حملات را بهبود می‌بخشند. دستاوردهای مورد انتظار این پژوهش شامل ارائه روشی نوین برای شناسایی حملات تزریق پایگاه‌داده غیررابطه‌ای، ایجاد مجموعه‌دادگان جامع و متنوع، و بهبود دقت مدل‌های یادگیری ماشینی در محیط‌های واقعی است.""")
    
    print("  ✓ Section 1-2 rewritten as flowing prose")

# مرحله 3: بازنویسی بخش 1-6 (نوآوری پژوهش)
print("\nStep 3: Rewriting section 1-6 (Innovation)...")

if '1-6' in sections:
    sec_1_6_start = sections['1-6']
    sec_1_6_end = sections.get('1-7', ch1_end)
    
    # حذف محتوای قدیمی (بعد از عنوان)
    for i in range(sec_1_6_start + 1, min(sec_1_6_end, len(doc.paragraphs))):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            if para.text.strip():
                para._element.getparent().remove(para._element)
    
    # پیدا کردن پاراگراف بعدی برای درج
    next_para_idx = sec_1_6_start + 1
    while next_para_idx < len(doc.paragraphs) and not doc.paragraphs[next_para_idx].text.strip():
        next_para_idx += 1
    
    # نوشتن متن جدید خلاصه و متمرکز
    new_para = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para.add_run("""نوآوری اصلی این پژوهش در تمرکز بر حملات تزریق پایگاه‌داده غیررابطه‌ای است که به‌دلیل تفاوت‌های ساختاری اساسی با پایگاه‌داده‌های رابطه‌ای، نیازمند رویکردهای امنیتی متفاوت و پیچیده‌تری هستند. برخلاف پایگاه‌داده‌های رابطه‌ای که از ساختار جدولی و زبان پرس‌وجوی استاندارد SQL استفاده می‌کنند، پایگاه‌داده‌های غیررابطه‌ای مانند MongoDB از ساختارهای غیرجدولی و زبان‌های پرس‌وجوی خاص خود بهره می‌برند. این تفاوت‌ها باعث می‌شود که حملات تزریق در این نوع پایگاه‌ها از الگوهای پیچیده‌تر و غیرمعمول‌تری استفاده کنند که شناسایی آن‌ها با روش‌های سنتی مبتنی بر SQL دشوار است.""")
    
    new_para2 = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para2.add_run("""در این پژوهش، با استفاده از مدل‌های زبان بزرگ برای تولید داده‌های مصنوعی مرتبط با حملات MongoDB، یک مجموعه‌داده جامع و متنوع ایجاد شده است. همچنین، یک سیستم یادگیری گروهی طراحی شده که از ترکیب چندین مدل یادگیری ماشینی برای بهبود دقت تشخیص استفاده می‌کند. این رویکرد جامع، خلأ موجود در پژوهش‌های امنیت سایبری را پر کرده و گامی مهم در راستای ارتقای امنیت پایگاه‌داده‌های غیررابطه‌ای برداشته است.""")
    
    print("  ✓ Section 1-6 rewritten and focused")

# مرحله 4: بازنویسی بخش 1-7 (ساختار فصول) به پاراگراف‌های پیوسته
print("\nStep 4: Rewriting section 1-7 (Chapter structure) as paragraphs...")

if '1-7' in sections:
    sec_1_7_start = sections['1-7']
    sec_1_7_end = ch1_end
    
    # خواندن محتوای فعلی
    current_content = []
    for i in range(sec_1_7_start, min(sec_1_7_end, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.text.strip() and 'فصل' in para.text:
            current_content.append(para.text.strip())
    
    # حذف محتوای قدیمی (بعد از عنوان)
    for i in range(sec_1_7_start + 1, min(sec_1_7_end, len(doc.paragraphs))):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            if para.text.strip():
                para._element.getparent().remove(para._element)
    
    # پیدا کردن پاراگراف بعدی برای درج
    next_para_idx = sec_1_7_start + 1
    while next_para_idx < len(doc.paragraphs) and not doc.paragraphs[next_para_idx].text.strip():
        next_para_idx += 1
    
    # نوشتن پاراگراف‌های جدید برای هر فصل
    new_para1 = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para1.add_run("""این پایان‌نامه در شش فصل تنظیم شده است و مسیر پژوهش را از مبانی نظری تا جمع‌بندی و پیشنهادهای آتی به‌صورت منسجم دنبال می‌کند.""")
    
    new_para2 = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para2.add_run("""فصل دوم به بررسی مبانی نظری و پیشینه پژوهش اختصاص دارد. در این فصل، مفاهیم پایه مرتبط با حملات تزریق در پایگاه‌داده‌های غیررابطه‌ای، کاربردهای یادگیری ماشین در امنیت، و رویکردهای تولید داده مصنوعی با مدل‌های زبان بزرگ در مقایسه با روش‌های سنتی مرور می‌شوند. همچنین، شکاف‌های پژوهشی و تمایزات رویکرد پیشنهادی جمع‌بندی می‌گردد.""")
    
    new_para3 = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para3.add_run("""فصل سوم روش پیشنهادی را تشریح می‌کند. این فصل شامل چارچوب کلی پژوهش، فرایند تولید داده مصنوعی مبتنی بر مدل‌های زبان بزرگ، معرفی مدل‌های یادگیری ماشین و مدل‌های گروهی به‌کاررفته و منطق انتخاب آن‌ها، تنظیمات کلیدی و پارامترها، و معماری پیاده‌سازی سیستم پیشنهادی است.""")
    
    new_para4 = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para4.add_run("""فصل چهارم به تحلیل نتایج و بحث اختصاص دارد. در این فصل، مراحل آماده‌سازی داده‌ها، اجرای مدل‌ها و نحوه ارزیابی بیان می‌شود. سپس عملکرد مدل‌ها با معیارهایی مانند دقت، شاخص‌های خطا، و معیارهای زمانی و کارایی مقایسه و تحلیل می‌گردد. همچنین، تفسیر یافته‌ها، مقایسه با پیشینه پژوهش، بررسی نقاط قوت و محدودیت‌ها، و تحلیل اثر داده‌های مصنوعی و انتخاب مدل‌ها بر عملکرد و تعمیم‌پذیری ارائه می‌شود.""")
    
    new_para5 = doc.paragraphs[next_para_idx].insert_paragraph_before()
    new_para5.add_run("""فصل پنجم شامل جمع‌بندی و کارهای آتی است. در این فصل، دستاوردهای اصلی پژوهش، مرور نوآوری‌ها، بیان محدودیت‌ها و ارائه پیشنهادهای مشخص برای توسعه و پژوهش‌های آینده ارائه می‌شود.""")
    
    print("  ✓ Section 1-7 rewritten as flowing paragraphs")

# مرحله 5: اصلاح ایرادات نگارشی
print("\nStep 5: Fixing grammatical and stylistic issues...")

# اصلاحات کلی
fixes = [
    ('پایگاه‌داده پایگاه‌داده', 'پایگاه‌داده'),
    ('مدل‌های زبان بزرگ مدل‌های زبان بزرگ', 'مدل‌های زبان بزرگ'),
    ('پایگاه‌داده رابطه‌ای پایگاه‌داده رابطه‌ای', 'پایگاه‌داده رابطه‌ای'),
    ('یادگیری ماشینی  ', 'یادگیری ماشینی '),
    ('یادگیری گروهی  ', 'یادگیری گروهی '),
]

for i in range(ch1_start, min(ch1_end, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text
    for old, new in fixes:
        if old in text:
            para.text = para.text.replace(old, new)

print("  ✓ Grammatical fixes applied")

# ذخیره فایل
doc.save(r'd:\Workingdir\Injection\Thesis_V3.4.docx')
print("\n✓ Chapter 1 refinement completed successfully!")
print("✓ File saved as Thesis_V3.4.docx")
