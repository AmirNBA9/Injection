from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import etree

def find_section_start(doc, heading_text):
    """پیدا کردن شروع یک بخش بر اساس عنوان"""
    for i, para in enumerate(doc.paragraphs):
        if heading_text in para.text and 'Heading' in para.style.name:
            return i
    return None

def find_section_end(doc, start_idx, next_heading_text):
    """پیدا کردن پایان یک بخش"""
    for i in range(start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if next_heading_text in para.text and 'Heading' in para.style.name:
            return i
    return len(doc.paragraphs)

# باز کردن سند
doc = Document(r'd:\Workingdir\Injection\Thesis_V3.3.docx')

print("Starting content transfer...")

# پیدا کردن موقعیت‌های کلیدی
print("\nFinding key positions...")

# پیدا کردن موقعیت‌های مقصد
dest_2_1_1_end = None  # انتهای 2-1-1
dest_2_1_2_end = None  # انتهای 2-1-2
dest_2_4_end = None    # انتهای 2-4
dest_ch3_start = None  # ابتدای فصل ۳

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if '2-1-1' in text and 'Heading' in style:
        # پیدا کردن انتهای 2-1-1 (قبل از 2-1-2)
        for j in range(i + 1, len(doc.paragraphs)):
            if '2-1-2' in doc.paragraphs[j].text and 'Heading' in doc.paragraphs[j].style.name:
                dest_2_1_1_end = j
                break
    
    if '2-1-2' in text and 'Heading' in style:
        # پیدا کردن انتهای 2-1-2 (قبل از 2-1-3)
        for j in range(i + 1, len(doc.paragraphs)):
            if '2-1-3' in doc.paragraphs[j].text and 'Heading' in doc.paragraphs[j].style.name:
                dest_2_1_2_end = j
                break
    
    if '2-4' in text and 'Heading' in style:
        # پیدا کردن انتهای 2-4 (قبل از فصل ۳)
        for j in range(i + 1, len(doc.paragraphs)):
            if 'فصل سوم' in doc.paragraphs[j].text or 'فصل ۳' in doc.paragraphs[j].text:
                dest_2_4_end = j
                break
    
    if 'فصل سوم' in text or 'فصل ۳' in text:
        # پیدا کردن ابتدای فصل ۳ (بعد از مقدمه، قبل از 3-1)
        for j in range(i + 1, min(i + 50, len(doc.paragraphs))):
            if '3-1' in doc.paragraphs[j].text and 'Heading' in doc.paragraphs[j].style.name:
                dest_ch3_start = j
                break

print(f"  - End of 2-1-1: {dest_2_1_1_end}")
print(f"  - End of 2-1-2: {dest_2_1_2_end}")
print(f"  - End of 2-4: {dest_2_4_end}")
print(f"  - Start of Chapter 3: {dest_ch3_start}")

# تعریف انتقال‌ها: (src_start, src_end, dest_idx, description)
transfers = []

if dest_2_1_1_end:
    transfers.append((192, 196, dest_2_1_1_end, "1-2-2 → end of 2-1-1"))
    transfers.append((197, 200, dest_2_1_1_end, "1-2-3 → end of 2-1-1"))

if dest_2_1_2_end:
    transfers.append((177, 191, dest_2_1_2_end, "1-2-1 → end of 2-1-2"))
    transfers.append((201, 224, dest_2_1_2_end, "1-2-4 → end of 2-1-2"))

if dest_ch3_start:
    transfers.append((225, 256, dest_ch3_start, "1-2-5 → beginning of Chapter 3"))

if dest_2_4_end:
    transfers.append((276, 296, dest_2_4_end, "1-4 → end of 2-4"))

# انجام انتقال‌ها از آخر به اول
print("\nExecuting transfers (from end to beginning)...")

for src_start, src_end, dest_idx, desc in reversed(transfers):
    print(f"\nTransferring: {desc} (P{src_start}-P{src_end}) → P{dest_idx}...")
    
    if src_start >= len(doc.paragraphs) or src_end >= len(doc.paragraphs):
        print(f"   ✗ Source range out of bounds")
        continue
    
    if dest_idx is None or dest_idx >= len(doc.paragraphs):
        print(f"   ✗ Destination index invalid")
        continue
    
    # کپی پاراگراف‌ها (از XML)
    body = doc._body._body
    source_elements = []
    
    for i in range(src_start, src_end + 1):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            source_elements.append(para._element)
    
    if not source_elements:
        print(f"   ✗ No paragraphs found")
        continue
    
    # پیدا کردن element مقصد
    dest_para = doc.paragraphs[dest_idx]
    dest_parent = dest_para._element.getparent()
    dest_element = dest_para._element
    
    # اضافه کردن علامت
    marker_para = doc.paragraphs[dest_idx].insert_paragraph_before()
    marker_run = marker_para.add_run("[منتقل‌شده از فصل ۱]")
    marker_run.bold = True
    marker_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # درج پاراگراف‌های کپی شده
    for elem in reversed(source_elements):  # معکوس برای حفظ ترتیب
        dest_parent.insert_before(elem, dest_element)
    
    print(f"   ✓ Transferred {len(source_elements)} paragraphs")

# حذف از مبدا (از آخر به اول)
print("\nRemoving original paragraphs from Chapter 1...")

# جمع‌آوری همه ایندکس‌های حذف
indices_to_remove = set()
for src_start, src_end, _, _ in transfers:
    for i in range(src_start, src_end + 1):
        indices_to_remove.add(i)

# حذف از آخر به اول
for i in sorted(indices_to_remove, reverse=True):
    if i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        para._element.getparent().remove(para._element)

print(f"   ✓ Removed {len(indices_to_remove)} paragraphs")

# ذخیره فایل
doc.save(r'd:\Workingdir\Injection\Thesis_V3.3.docx')
print("\n✓ All transfers completed successfully!")
print("✓ File saved as Thesis_V3.3.docx")
