from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document(r'd:\Workingdir\Injection\Thesis_V3.4.docx')

print("Analyzing Chapter 1...")

# پیدا کردن محدوده فصل اول
ch1_start = None
ch1_end = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if 'فصل اول' in text and 'Heading' in style:
        ch1_start = i
    if 'فصل دوم' in text and 'Heading' in style:
        ch1_end = i
        break

print(f"Chapter 1: P{ch1_start} to P{ch1_end}")

# پیدا کردن بخش‌های مختلف
sections = {}
for i in range(ch1_start, min(ch1_end, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if 'Heading' in style:
        if '1-1' in text:
            sections['1-1'] = i
        elif '1-2' in text:
            sections['1-2'] = i
        elif '1-3' in text:
            sections['1-3'] = i
        elif '1-5' in text:
            sections['1-5'] = i
        elif '1-6' in text:
            sections['1-6'] = i
        elif '1-7' in text:
            sections['1-7'] = i

print(f"\nFound sections: {sections}")

# بررسی محتوای بخش 1-2
print("\nSection 1-2 content preview:")
if '1-2' in sections:
    start = sections['1-2']
    end = sections.get('1-3', ch1_end)
    for i in range(start, min(start + 30, end, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.text.strip():
            print(f"P{i}: {para.text[:100]}")

print("\n✓ Analysis complete. Ready for refinement.")
