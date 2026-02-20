from docx import Document

doc = Document(r'd:\Workingdir\Injection\Thesis_V3.3.docx')

print("Verifying transfers...\n")

# پیدا کردن علامت‌های انتقال
markers = []
for i, para in enumerate(doc.paragraphs):
    if '[منتقل' in para.text:
        markers.append(i)

print(f"Found {len(markers)} transfer markers:")
for idx in markers:
    if idx + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[idx + 1]
        print(f"  P{idx}: Next para starts with: {next_para.text[:70]}")

# بررسی اینکه آیا بخش‌های حذف شده از فصل ۱ حذف شده‌اند
print("\nChecking Chapter 1 for removed sections...")
ch1_start = None
ch1_end = None

for i, para in enumerate(doc.paragraphs):
    if 'فصل اول' in para.text and 'Heading' in (para.style.name if para.style else ''):
        ch1_start = i
    if 'فصل دوم' in para.text and 'Heading' in (para.style.name if para.style else ''):
        ch1_end = i
        break

if ch1_start and ch1_end:
    removed_sections = ['1-2-1', '1-2-2', '1-2-3', '1-2-4', '1-2-5', '1-4']
    found_in_ch1 = []
    
    for i in range(ch1_start, min(ch1_end, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if any(sec in para.text for sec in removed_sections) and 'Heading' in (para.style.name if para.style else ''):
            found_in_ch1.append(f"P{i}: {para.text[:50]}")
    
    if found_in_ch1:
        print(f"  ✗ Found {len(found_in_ch1)} removed sections still in Chapter 1:")
        for f in found_in_ch1:
            print(f"    {f}")
    else:
        print("  ✓ All removed sections successfully deleted from Chapter 1")

print("\n✓ Verification complete!")
