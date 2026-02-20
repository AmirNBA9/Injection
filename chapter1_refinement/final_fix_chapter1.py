from docx import Document

doc = Document(r'd:\Workingdir\Injection\Thesis_V3.4.docx')

print("Final fixes for Chapter 1...\n")

# پیدا کردن بخش 1-2
sec_1_2_idx = None
for i, para in enumerate(doc.paragraphs):
    if '1-2' in para.text and '1-2-' not in para.text and 'Heading' in (para.style.name if para.style else ''):
        sec_1_2_idx = i
        break

if sec_1_2_idx:
    print(f"Found section 1-2 at P{sec_1_2_idx}")
    
    # خواندن پاراگراف‌های فعلی
    current_paras = []
    for i in range(sec_1_2_idx + 1, min(sec_1_2_idx + 10, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text and len(text) > 50:
            current_paras.append((i, text))
            if len(current_paras) >= 3:
                break
    
    print(f"Current paragraphs after 1-2:")
    for idx, text in current_paras:
        print(f"  P{idx}: {text[:100]}")
    
    # اگر ترتیب معکوس است، اصلاح کنیم
    if len(current_paras) >= 3:
        # بررسی ترتیب: باید با "این پژوهش با هدف" شروع شود
        first_para = current_paras[0][1]
        if 'این پژوهش شامل طراحی' in first_para:
            print("\n  Fixing order: paragraphs are reversed")
            
            # حذف پاراگراف‌های فعلی
            for idx, _ in current_paras:
                if idx < len(doc.paragraphs):
                    para = doc.paragraphs[idx]
                    para._element.getparent().remove(para._element)
            
            # پیدا کردن نقطه درج
            next_idx = sec_1_2_idx + 1
            while next_idx < len(doc.paragraphs) and not doc.paragraphs[next_idx].text.strip():
                next_idx += 1
            
            # نوشتن به ترتیب صحیح
            para1 = doc.paragraphs[next_idx].insert_paragraph_before()
            para1.add_run("""این پژوهش با هدف ارائه راهکاری جامع برای تشخیص حملات تزریق پایگاه‌داده غیررابطه‌ای طراحی شده است. برای دست‌یابی به این هدف، مجموعه‌داده‌های اولیه به‌صورت دستی تحلیل و برچسب‌گذاری می‌شوند تا کیفیت داده‌های پایه تضمین گردد. سپس از مدل‌های زبان بزرگ برای تولید داده‌های مصنوعی متنوع و واقعی‌تر استفاده می‌شود که به غنی‌سازی مجموعه‌دادگان کمک می‌کند. در مرحله بعد، تکنیک‌های مختلف یادگیری ماشینی برای ارزیابی و دسته‌بندی داده‌های موجود و مصنوعی به‌کار گرفته می‌شوند.""")
            
            para2 = doc.paragraphs[next_idx].insert_paragraph_before()
            para2.add_run("""مفاهیم کلیدی این پژوهش شامل شناسایی و مقابله با حملات تزریق پایگاه‌داده غیررابطه‌ای است که از نقاط ضعف این نوع پایگاه‌ها سوءاستفاده می‌کنند. این حملات معمولاً از طریق ارسال درخواست‌های مخرب و دست‌کاری پرس‌وجوها انجام می‌گیرند. همچنین، تحلیل و طراحی مجموعه‌دادگان شامل داده‌های اولیه و داده‌های مصنوعی تولیدشده با مدل‌های زبان بزرگ، یکی از محورهای اصلی پژوهش است. در این راستا، به‌جای استفاده از شبکه‌های مولد متخاصم، از مدل‌های زبان بهره گرفته می‌شود که توانایی تولید داده‌های متنی و ساختاریافته باکیفیت بالا را دارند.""")
            
            para3 = doc.paragraphs[next_idx].insert_paragraph_before()
            para3.add_run("""این پژوهش شامل طراحی و ارزیابی یک سیستم یادگیری گروهی است که از مدل‌های مختلف مانند رگرسیون لجستیک، جنگل تصادفی و XGBoost تشکیل شده است. این مدل‌ها با ترکیب قدرت یادگیری فردی خود، دقت شناسایی حملات را بهبود می‌بخشند. دستاوردهای مورد انتظار این پژوهش شامل ارائه روشی نوین برای شناسایی حملات تزریق پایگاه‌داده غیررابطه‌ای، ایجاد مجموعه‌دادگان جامع و متنوع، و بهبود دقت مدل‌های یادگیری ماشینی در محیط‌های واقعی است.""")
            
            print("  ✓ Order fixed")

# بررسی وجود بخش 1-7
print("\nChecking for section 1-7...")
found_1_7 = False
for i, para in enumerate(doc.paragraphs):
    if '1-7' in para.text and 'Heading' in (para.style.name if para.style else ''):
        found_1_7 = True
        print(f"  Found section 1-7 at P{i}")
        
        # بررسی محتوا
        for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            if p.text.strip():
                print(f"    P{j}: {p.text[:100]}")
        break

if not found_1_7:
    print("  Section 1-7 not found - may have been removed by user")

# ذخیره
doc.save(r'd:\Workingdir\Injection\Thesis_V3.4.docx')
print("\n✓ Final fixes applied!")
print("✓ File saved as Thesis_V3.4.docx")
