from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# باز کردن سند
doc = Document(r'd:\Workingdir\Injection\Thesis_V3.3.docx')

print("Starting content transfer...")

# پیدا کردن موقعیت‌های کلیدی
print("\nFinding key positions...")

dest_2_1_1_end = None
dest_2_1_2_end = None
dest_2_4_end = None
dest_ch3_start = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if '2-1-1' in text and 'Heading' in style:
        for j in range(i + 1, len(doc.paragraphs)):
            if j < len(doc.paragraphs):
                p = doc.paragraphs[j]
                if '2-1-2' in p.text and 'Heading' in (p.style.name if p.style else ''):
                    dest_2_1_1_end = j
                    break
    
    if '2-1-2' in text and 'Heading' in style:
        for j in range(i + 1, len(doc.paragraphs)):
            if j < len(doc.paragraphs):
                p = doc.paragraphs[j]
                if '2-1-3' in p.text and 'Heading' in (p.style.name if p.style else ''):
                    dest_2_1_2_end = j
                    break
    
    if '2-4' in text and 'Heading' in style:
        for j in range(i + 1, len(doc.paragraphs)):
            if j < len(doc.paragraphs):
                p = doc.paragraphs[j]
                if 'فصل سوم' in p.text or 'فصل ۳' in p.text:
                    dest_2_4_end = j
                    break
    
    if 'فصل سوم' in text or 'فصل ۳' in text:
        for j in range(i + 1, min(i + 50, len(doc.paragraphs))):
            if j < len(doc.paragraphs):
                p = doc.paragraphs[j]
                if '3-1' in p.text and 'Heading' in (p.style.name if p.style else ''):
                    dest_ch3_start = j
                    break

print(f"  - End of 2-1-1: {dest_2_1_1_end}")
print(f"  - End of 2-1-2: {dest_2_1_2_end}")
print(f"  - End of 2-4: {dest_2_4_end}")
print(f"  - Start of Chapter 3: {dest_ch3_start}")

def transfer_section(src_start, src_end, dest_idx, desc):
    """انتقال یک بخش با کپی محتوا"""
    print(f"\nTransferring: {desc} (P{src_start}-P{src_end}) → P{dest_idx}...")
    
    if src_start >= len(doc.paragraphs) or src_end >= len(doc.paragraphs):
        print(f"   ✗ Source range out of bounds")
        return False
    
    if dest_idx is None or dest_idx >= len(doc.paragraphs):
        print(f"   ✗ Destination index invalid")
        return False
    
    # کپی محتوای پاراگراف‌ها
    copied_data = []
    for i in range(src_start, src_end + 1):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            data = {
                'text': para.text,
                'style': para.style.name if para.style else None,
                'runs': []
            }
            # کپی formatting از runs
            for run in para.runs:
                data['runs'].append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic
                })
            copied_data.append(data)
    
    if not copied_data:
        print(f"   ✗ No paragraphs found")
        return False
    
    # درج علامت
    marker_para = doc.paragraphs[dest_idx].insert_paragraph_before()
    marker_run = marker_para.add_run("[منتقل‌شده از فصل ۱]")
    marker_run.bold = True
    marker_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # درج پاراگراف‌های کپی شده
    for data in copied_data:
        new_para = doc.paragraphs[dest_idx].insert_paragraph_before()
        new_para.text = data['text']
        if data['style']:
            try:
                new_para.style = data['style']
            except:
                pass
        
        # بازسازی formatting
        if data['runs']:
            new_para.clear()
            for run_data in data['runs']:
                run = new_para.add_run(run_data['text'])
                if run_data.get('bold'):
                    run.bold = True
                if run_data.get('italic'):
                    run.italic = True
    
    print(f"   ✓ Transferred {len(copied_data)} paragraphs")
    return True

# انجام انتقال‌ها از آخر به اول
print("\nExecuting transfers (from end to beginning)...")

transfers_done = []

if dest_2_1_1_end:
    if transfer_section(197, 200, dest_2_1_1_end, "1-2-3 → end of 2-1-1"):
        transfers_done.append((197, 200))
    if transfer_section(192, 196, dest_2_1_1_end, "1-2-2 → end of 2-1-1"):
        transfers_done.append((192, 196))

if dest_2_1_2_end:
    if transfer_section(201, 224, dest_2_1_2_end, "1-2-4 → end of 2-1-2"):
        transfers_done.append((201, 224))
    if transfer_section(177, 191, dest_2_1_2_end, "1-2-1 → end of 2-1-2"):
        transfers_done.append((177, 191))

if dest_ch3_start:
    if transfer_section(225, 256, dest_ch3_start, "1-2-5 → beginning of Chapter 3"):
        transfers_done.append((225, 256))

if dest_2_4_end:
    if transfer_section(276, 296, dest_2_4_end, "1-4 → end of 2-4"):
        transfers_done.append((276, 296))

# حذف از مبدا (از آخر به اول)
print("\nRemoving original paragraphs from Chapter 1...")

indices_to_remove = set()
for src_start, src_end in transfers_done:
    for i in range(src_start, src_end + 1):
        indices_to_remove.add(i)

# حذف از آخر به اول
removed_count = 0
for i in sorted(indices_to_remove, reverse=True):
    if i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        para._element.getparent().remove(para._element)
        removed_count += 1

print(f"   ✓ Removed {removed_count} paragraphs")

# ذخیره فایل
doc.save(r'd:\Workingdir\Injection\Thesis_V3.3.docx')
print("\n✓ All transfers completed successfully!")
print("✓ File saved as Thesis_V3.3.docx")
