from docx import Document

doc = Document(r'd:\Workingdir\Injection\Thesis_V3.4.docx')

print("Verifying Chapter 1 refinements...\n")

# پیدا کردن فصل اول
ch1_start = None
for i, para in enumerate(doc.paragraphs):
    if 'فصل اول' in para.text and 'Heading' in (para.style.name if para.style else ''):
        ch1_start = i
        break

print(f"Chapter 1 starts at: P{ch1_start}")

# بررسی بخش 1-2
print("\n=== Section 1-2 (should be flowing prose) ===")
found_1_2 = False
for i in range(ch1_start, min(ch1_start + 50, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if '1-2' in text and 'Heading' in (para.style.name if para.style else ''):
        found_1_2 = True
        print(f"P{i}: {text}")
        # نمایش چند پاراگراف بعد
        for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            if p.text.strip():
                print(f"  P{j}: {p.text[:150]}")

# بررسی بخش 1-6
print("\n=== Section 1-6 (should be concise and focused) ===")
for i in range(ch1_start, min(ch1_start + 300, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if '1-6' in text and 'Heading' in (para.style.name if para.style else ''):
        print(f"P{i}: {text}")
        # نمایش چند پاراگراف بعد
        for j in range(i + 1, min(i + 3, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            if p.text.strip():
                print(f"  P{j}: {p.text[:150]}")

# بررسی بخش 1-7
print("\n=== Section 1-7 (should be paragraphs, not bullets) ===")
for i in range(ch1_start, min(ch1_start + 300, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if '1-7' in text and 'Heading' in (para.style.name if para.style else ''):
        print(f"P{i}: {text}")
        # نمایش چند پاراگراف بعد
        for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            if p.text.strip():
                print(f"  P{j}: {p.text[:150]}")

print("\n✓ Verification complete!")
