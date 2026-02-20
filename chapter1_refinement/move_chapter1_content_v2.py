from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy

def find_paragraph_by_text(doc, search_text, start_idx=0, end_idx=None):
    """پیدا کردن پاراگراف بر اساس متن"""
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    for i in range(start_idx, end_idx):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            if search_text in para.text:
                return i
    return None

def copy_paragraph_element(para):
    """کپی کامل element پاراگراف"""
    return deepcopy(para._element)

def insert_marker_paragraph(doc, target_idx):
    """درج پاراگراف علامت‌گذاری"""
    marker_para = doc.paragraphs[target_idx].insert_paragraph_before()
    marker_run = marker_para.add_run("[منتقل‌شده از فصل ۱]")
    marker_run.bold = True
    marker_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    return target_idx + 1

# باز کردن سند
doc = Document(r'd:\Workingdir\Injection\Thesis_V3.3.docx')

print("Starting content transfer (from end to beginning to preserve indices)...")

# تعریف همه انتقال‌ها: (src_start, src_end, dest_marker, dest_after)
transfers = [
    # (src_start, src_end, dest_search_text, dest_after, description)
    (276, 296, 'فصل سوم', False, "1-4 → end of 2-4"),
    (225, 256, '3-1', False, "1-2-5 → beginning of Chapter 3"),
    (201, 224, '2-1-3', False, "1-2-4 → end of 2-1-2"),
    (177, 191, '2-1-3', False, "1-2-1 → end of 2-1-2"),
    (197, 200, '2-1-2', False, "1-2-3 → end of 2-1-1"),
    (192, 196, '2-1-2', False, "1-2-2 → end of 2-1-1"),
]

# انجام انتقال‌ها از آخر به اول
for src_start, src_end, dest_marker, dest_after, desc in transfers:
    print(f"\nTransferring: {desc} (P{src_start}-P{src_end})...")
    
    # پیدا کردن مقصد
    dest_idx = None
    if dest_marker == 'فصل سوم':
        # برای فصل سوم، پیدا کردن عنوان فصل
        for i in range(500, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if 'فصل سوم' in para.text or 'فصل ۳' in para.text:
                dest_idx = i
                break
    elif dest_marker == '3-1':
        # برای فصل ۳، پیدا کردن 3-1
        for i in range(800, 900):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                if '3-1' in para.text and 'Heading' in para.style.name:
                    dest_idx = i
                    break
    elif dest_marker == '2-1-3':
        # برای 2-1-2، پیدا کردن 2-1-3
        for i in range(400, 500):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                if '2-1-3' in para.text and 'Heading' in para.style.name:
                    dest_idx = i
                    break
    elif dest_marker == '2-1-2':
        # برای 2-1-1، پیدا کردن 2-1-2
        for i in range(400, 500):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                if '2-1-2' in para.text and 'Heading' in para.style.name:
                    dest_idx = i
                    break
    
    if dest_idx is None:
        print(f"   ✗ Could not find destination marker '{dest_marker}'")
        continue
    
    # کپی پاراگراف‌ها
    copied_elements = []
    for i in range(src_start, src_end + 1):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            copied_elements.append(copy_paragraph_element(para))
    
    if not copied_elements:
        print(f"   ✗ No paragraphs found in range {src_start}-{src_end}")
        continue
    
    # درج در مقصد
    if dest_idx < len(doc.paragraphs):
        # اضافه کردن علامت
        dest_idx = insert_marker_paragraph(doc, dest_idx)
        
        # درج پاراگراف‌های کپی شده
        for elem in copied_elements:
            doc.paragraphs[dest_idx]._element.getparent().insert_before(
                elem, doc.paragraphs[dest_idx]._element
            )
            dest_idx += 1
        
        print(f"   ✓ Transferred {len(copied_elements)} paragraphs to position {dest_idx - len(copied_elements)}")
    else:
        print(f"   ✗ Destination index {dest_idx} out of range")
        continue

# حالا حذف از مبدا (از آخر به اول)
print("\nRemoving original paragraphs from Chapter 1...")
for src_start, src_end, _, _, desc in reversed(transfers):
    print(f"Removing: {desc} (P{src_start}-P{src_end})...")
    # حذف از آخر به اول
    for i in range(src_end, src_start - 1, -1):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            para._element.getparent().remove(para._element)
    print(f"   ✓ Removed paragraphs {src_start}-{src_end}")

# ذخیره فایل
doc.save(r'd:\Workingdir\Injection\Thesis_V3.3.docx')
print("\n✓ All transfers completed successfully!")
print("✓ File saved as Thesis_V3.3.docx")
